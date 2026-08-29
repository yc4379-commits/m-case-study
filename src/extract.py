"""
Document extraction layer.

This module turns a resume file (.docx / .pdf) into clean text plus a quality
report. It is deliberately separate from the LLM parsing step, because the
single biggest source of error in a resume pipeline is not the model -- it is
losing content before the model ever sees it.

Three failure modes are handled explicitly, each observed in the sample data:

  1. Word tables. python-docx's `Document.paragraphs` does NOT include text
     inside tables. Several resumes in this corpus place their entire work
     history in tables, so naive paragraph extraction silently drops most of
     the document and the downstream LLM confidently parses a fragment.

  2. PDF multi-column layout. A two-column CV read left-to-right interleaves
     the sidebar (skills, languages) into the middle of work-experience lines,
     destroying sentence structure.

  3. Ligature / encoding corruption. Some PDF producers emit `ti` and `ffi`
     ligatures that decode to U+FFFD, turning "Quantitative" into "Quan?ta?ve".

Every extraction returns an ExtractionReport so that downstream code -- and the
end user -- can see how much to trust the result. We never silently degrade.
"""

from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Any

import docx
import pdfplumber


# --------------------------------------------------------------------------
# Quality report
# --------------------------------------------------------------------------

@dataclass
class ExtractionReport:
    """Diagnostics describing how well we read a document.

    These fields are a by-product of extraction that most pipelines throw
    away. We keep them because they drive parse-confidence downstream.
    """

    source_file: str
    file_type: str
    char_count: int = 0
    paragraph_chars: int = 0
    table_chars: int = 0
    table_count: int = 0
    textbox_chars: int = 0
    textbox_count: int = 0
    page_count: int = 0
    multi_column_detected: bool = False
    ligature_repairs: int = 0
    replacement_chars_remaining: int = 0
    warnings: list[str] = field(default_factory=list)

    @property
    def table_share(self) -> float:
        """Fraction of extracted text that came from tables."""
        return self.table_chars / self.char_count if self.char_count else 0.0

    def to_dict(self) -> dict[str, Any]:
        d = asdict(self)
        d["table_share"] = round(self.table_share, 3)
        return d


# --------------------------------------------------------------------------
# Text repair
# --------------------------------------------------------------------------

# Ligature code points that survive in well-formed PDFs. We normalise these to
# their ASCII expansions so the model sees ordinary words.
_LIGATURES = {
    "ﬀ": "ff", "ﬁ": "fi", "ﬂ": "fl",
    "ﬃ": "ffi", "ﬄ": "ffl", "ﬅ": "st", "ﬆ": "st",
}

# When a PDF's ToUnicode map is broken, ligature glyphs decode to U+FFFD
# instead. The character itself is lost, but the surrounding letters constrain
# what it must have been.
#
# These rules are derived from the observed corpus rather than guessed. Every
# damaged token in the sample PDFs was enumerated before writing them: of 26
# occurrences, 25 are "ti" (Educa?on, Stochas?c, deriva?ves, Sta?s?cs, ...)
# and one is "tf" (Por?olio). An earlier version added an "ffi" rule for the
# `?c` context, which then corrupted "Sta?s?cs" into "Statisffics" -- the
# model reported it in extraction_notes, which is precisely what those notes
# are for.
#
# The generic rule is therefore "ti", ordered last so specific contexts win.
# A production system with a broader corpus would validate candidates against
# a lexicon instead of defaulting; residual damage is reported either way.
_FFFD_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"(?<=[Pp]or)�(?=olio)"), "tf"),          # Por?olio -> Portfolio
    (re.compile(r"(?<=[A-Za-z])�(?=[A-Za-z])"), "ti"),    # dominant case
]


def repair_text(raw: str) -> tuple[str, int, int]:
    """Normalise unicode and repair ligature damage.

    Returns (repaired_text, repairs_made, replacement_chars_remaining).

    We report what we could not fix rather than hiding it: a document with
    many surviving U+FFFD is one whose parse should be treated as suspect.
    """
    text = unicodedata.normalize("NFKC", raw)

    repairs = 0
    for lig, expansion in _LIGATURES.items():
        if lig in text:
            repairs += text.count(lig)
            text = text.replace(lig, expansion)

    for pattern, expansion in _FFFD_PATTERNS:
        text, n = pattern.subn(expansion, text)
        repairs += n

    remaining = text.count("�")

    # Collapse runs of spaces/tabs but preserve line structure, which carries
    # meaningful layout signal (job titles, dates, section headers).
    text = re.sub(r"[ \t ]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(line.strip() for line in text.split("\n"))

    return text.strip(), repairs, remaining


# --------------------------------------------------------------------------
# Word (.docx)
# --------------------------------------------------------------------------

def _dedupe_row(cells: list[str]) -> list[str]:
    """Drop empty and repeated cells from a table row.

    Merged cells in Word are reported once per underlying grid column, so a
    row spanning three columns yields the same string three times. Left in,
    this triples the token cost and biases the model toward repeated phrases.
    """
    out: list[str] = []
    for c in cells:
        c = c.strip()
        if c and (not out or c != out[-1]) and c not in out:
            out.append(c)
    return out


_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _textbox_text(element) -> list[str]:
    """Text inside floating text boxes anchored to this element.

    python-docx exposes neither `Document.paragraphs` nor `Document.tables`
    content that lives inside a text box: `Paragraph.text` concatenates only
    the paragraph's own direct runs, and text-box runs are nested under
    w:pict / w:drawing -> w:txbxContent.

    This is not an edge case. In this corpus one resume puts its candidate
    name, degree, and the "PROFESSIONAL EXPERIENCE" section heading in text
    boxes. Reading only paragraphs and tables, that document appears to have
    no name and an unlabelled work-history table -- and every downstream
    conclusion drawn from that appearance is wrong. The document was fine;
    the reader was not.
    """
    out: list[str] = []
    for box in element.iter(f"{_W_NS}txbxContent"):
        text = "".join(node.text or "" for node in box.iter(f"{_W_NS}t")).strip()
        if text and text not in out:
            out.append(text)
    return out


def extract_docx(path: Path) -> tuple[str, ExtractionReport]:
    """Extract paragraphs, tables AND text boxes from a .docx, in document order.

    Document order matters: a table that follows an "EXPERIENCE" heading is
    experience. Appending all tables at the end -- the common shortcut --
    severs that link and makes section attribution unrecoverable.
    """
    report = ExtractionReport(source_file=path.name, file_type="docx")
    document = docx.Document(str(path))

    # Walk the body's XML children so paragraphs and tables stay interleaved
    # exactly as they appear in the document.
    body = document.element.body
    para_map = {p._element: p for p in document.paragraphs}
    table_map = {t._element: t for t in document.tables}

    chunks: list[str] = []
    para_chars = 0
    table_chars = 0

    # Floating text boxes carry no position in the document flow: they are
    # anchored to whatever paragraph happens to sit near them, and their
    # visual placement lives in anchor geometry we do not compute. Emitting
    # them where they are anchored put this resume's name and its
    # "PROFESSIONAL EXPERIENCE" heading two-thirds of the way down the text,
    # which reads as the start of a new section rather than the top of the CV.
    #
    # In resumes, floating boxes almost always hold header material -- name,
    # degree, section banners. Hoisting them to the front is therefore closer
    # to true reading order than anchor order is. It remains an approximation,
    # so the count is reported and a warning is raised.
    boxes: list[str] = []
    for text in _textbox_text(body):
        if text not in boxes:
            boxes.append(text)
    if boxes:
        chunks.extend(boxes)
        report.textbox_chars = sum(len(b) for b in boxes)

    for child in body.iterchildren():
        if child in para_map:
            text = para_map[child].text.strip()
            if text:
                chunks.append(text)
                para_chars += len(text)
        elif child in table_map:
            report.table_count += 1
            for row in table_map[child].rows:
                cells = _dedupe_row([c.text for c in row.cells])
                if cells:
                    line = " | ".join(cells)
                    chunks.append(line)
                    table_chars += len(line)

    report.textbox_count = len(boxes)

    raw = "\n".join(chunks)
    text, repairs, remaining = repair_text(raw)

    report.char_count = len(text)
    report.paragraph_chars = para_chars
    report.table_chars = table_chars
    report.ligature_repairs = repairs
    report.replacement_chars_remaining = remaining

    if report.char_count < 1000:
        report.warnings.append(
            f"Very little text extracted ({report.char_count} chars) - "
            "document may be image-based or use unsupported structure."
        )
    if report.textbox_count:
        report.warnings.append(
            f"{report.textbox_count} floating text box(es) recovered and "
            "hoisted to the top of the text; paragraph-and-table extraction "
            "alone would have missed this content entirely."
        )
    if report.table_share > 0.5:
        report.warnings.append(
            f"{report.table_share:.0%} of content came from tables - "
            "naive paragraph-only extraction would have lost most of this file."
        )
    return text, report


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------

_LINE_TOLERANCE = 3.0  # points; words within this vertical distance share a line


def _group_lines(words: list[dict]) -> list[list[dict]]:
    """Group words into visual lines by vertical position."""
    lines: dict[int, list[dict]] = {}
    for w in words:
        lines.setdefault(round(w["top"] / _LINE_TOLERANCE), []).append(w)
    return [sorted(ws, key=lambda w: w["x0"]) for _, ws in sorted(lines.items())]


def _detect_columns(page) -> float | None:
    """Return an x-coordinate splitting a two-column page, or None.

    A first attempt used the histogram of *word* start positions, looking for
    an empty vertical gutter. That fails on real CVs: a wide main column
    produces word starts across most of the page, so no gutter exists even
    when the layout is plainly two-column.

    What actually separates the two cases is the distribution of *line* start
    positions. Main-column lines all begin at the left margin; sidebar lines
    all begin at the sidebar's margin. The result is strongly bimodal with a
    wide gap between the two clusters, regardless of how far main-column text
    extends to the right.

    Three guards keep this from firing on single-column pages:
      - the gap must be a substantial share of page width;
      - the right-hand cluster must contain several lines (a sidebar has many
        entries; a lone centred header line does not);
      - both sides must hold a meaningful share of the page's words.
    """
    words = page.extract_words()
    if len(words) < 40:
        return None

    width = page.width
    line_starts = sorted(min(w["x0"] for w in line) for line in _group_lines(words))

    # Widest gap between distinct line-start positions, with the right-hand
    # cluster sitting in the right half of the page where a sidebar lives.
    best_gap, right_min = 0.0, None
    unique = sorted(set(line_starts))
    for lower, upper in zip(unique, unique[1:]):
        if upper - lower > best_gap and 0.45 * width <= upper <= 0.95 * width:
            best_gap, right_min = upper - lower, upper

    if right_min is None or best_gap < 0.15 * width:
        return None

    split_x = right_min - 3.0  # sit just left of the sidebar margin

    right_lines = [s for s in line_starts if s >= split_x]
    if len(right_lines) < 4 or len(right_lines) < 0.15 * len(line_starts):
        return None

    left_words = sum(1 for w in words if w["x0"] < split_x)
    if min(left_words, len(words) - left_words) < 0.08 * len(words):
        return None

    return split_x


def _render_column(words: list[dict]) -> str:
    """Rebuild text for one column, reading top-to-bottom."""
    return "\n".join(
        " ".join(w["text"] for w in line) for line in _group_lines(words)
    )


def _extract_page(page) -> tuple[str, bool]:
    """Extract one PDF page, separating columns when present.

    Words are assigned to a column by their own start position rather than by
    cropping the page. Cropping would also capture main-column words that
    merely extend past the split, reintroducing the interleaving we are
    trying to remove.
    """
    split_x = _detect_columns(page)
    if split_x is None:
        return page.extract_text() or "", False

    words = page.extract_words()
    left = [w for w in words if w["x0"] < split_x]
    right = [w for w in words if w["x0"] >= split_x]
    return f"{_render_column(left)}\n{_render_column(right)}", True


def extract_pdf(path: Path) -> tuple[str, ExtractionReport]:
    """Extract text from a PDF, handling multi-column layouts."""
    report = ExtractionReport(source_file=path.name, file_type="pdf")
    pages: list[str] = []

    with pdfplumber.open(str(path)) as pdf:
        report.page_count = len(pdf.pages)
        for page in pdf.pages:
            text, was_split = _extract_page(page)
            if was_split:
                report.multi_column_detected = True
            pages.append(text)

    text, repairs, remaining = repair_text("\n".join(pages))

    report.char_count = len(text)
    report.paragraph_chars = len(text)
    report.ligature_repairs = repairs
    report.replacement_chars_remaining = remaining

    if report.multi_column_detected:
        report.warnings.append(
            "Multi-column layout detected - columns were separated before "
            "reading. Without this, sidebar content interleaves into the "
            "work-experience text."
        )
    if remaining:
        report.warnings.append(
            f"{remaining} unrecoverable character(s) remain after ligature "
            "repair - source PDF has a damaged character map."
        )
    if report.char_count < 1000:
        report.warnings.append(
            f"Very little text extracted ({report.char_count} chars) - "
            "PDF may be a scan requiring OCR."
        )
    return text, report


# --------------------------------------------------------------------------
# Entry point
# --------------------------------------------------------------------------

_EXTRACTORS = {".docx": extract_docx, ".pdf": extract_pdf}


def extract(path: str | Path) -> tuple[str, ExtractionReport]:
    """Extract clean text and a quality report from one resume file."""
    path = Path(path)
    extractor = _EXTRACTORS.get(path.suffix.lower())
    if extractor is None:
        raise ValueError(
            f"Unsupported file type '{path.suffix}'. Supported: "
            f"{', '.join(sorted(_EXTRACTORS))}"
        )
    return extractor(path)


def extract_directory(directory: str | Path) -> list[tuple[str, ExtractionReport]]:
    """Extract every supported document in a directory, sorted by filename.

    A file that fails to parse produces an empty result with the error
    recorded in its report, rather than aborting the batch. At scale one
    corrupt file must never stop the run.
    """
    directory = Path(directory)
    results = []
    for path in sorted(directory.iterdir()):
        if path.suffix.lower() not in _EXTRACTORS:
            continue
        try:
            results.append(extract(path))
        except Exception as exc:  # noqa: BLE001 - batch robustness is the point
            report = ExtractionReport(
                source_file=path.name,
                file_type=path.suffix.lstrip("."),
                warnings=[f"Extraction failed: {type(exc).__name__}: {exc}"],
            )
            results.append(("", report))
    return results
